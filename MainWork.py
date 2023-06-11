import tkinter.messagebox as mb
import random

from tkinter import *
from tkinter import ttk
from tkinter import filedialog as fd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX



class MainWork(Tk):
    def __init__(self):
        super().__init__()
        self.protocol("WM_DELETE_WINDOW", self.confirm_delete)
    
        self.path = ""
        self.dictQuestions = {}
        self.AnswerQuestions = {}
        self.NoAnswerQuestions = {}
        self.name = ""

        self.questNumber = 0

        self.title("")
        w = 600
        h = 300
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        x = (sw - w) / 2
        y = (sh - h) / 2
        self.geometry('%dx%d+%d+%d' % (w, h, x, y))
        self.resizable(width=False, height=False)

        
        self.mainmenu = Menu(self) 
        self.config(menu=self.mainmenu) 
        
        self.filemenu = Menu(self.mainmenu, tearoff=0)
        self.filemenu.add_command(label="Загрузить файл", command=self.file)
        self.filemenu.add_command(label=f"Сохранить результат", command=self.save_result)
        self.filemenu.add_command(label="Выход", command=self.confirm_delete)

        self.helpmenu = Menu(self.mainmenu, tearoff=0)
        self.helpmenu.add_command(label="Помощь")
        self.helpmenu.add_command(label="О программе")

        self.mainmenu.add_cascade(label="Файл",
                     menu=self.filemenu)
        self.mainmenu.add_cascade(label="Справка",
                     menu=self.helpmenu)   


        self.label = ttk.Label(self, text="Это всплывающее окно")
        self.button = ttk.Button(self, text="Закрыть", command=self.destroy)
        

        self.nameFile = ttk.Label(self, text=f'')
        self.nameFile.place(x=5, y=7)
        self.getQuestions = ttk.Button(self, text="Получить вопрос", command=self.getQuestion)
        self.getQuestions.place(x=235, y=210)
        self.nextQues = ttk.Button(self, text=">>>", command=self.getNextQuest, state='disabled')
        self.nextQues.place(x=350, y=210)
        self.beforeQues = ttk.Button(self, text="<<<", command=self.getBeforeQuest, state='disabled')
        self.beforeQues.place(x=150, y=210)

        self.var = IntVar()
        self.var.set(0)
        self.labelRadio = ttk.Label(self, text=f'Какие вопросы ты хочешь получать?', font="times 11 italic")
        self.rad4 = ttk.Checkbutton(self, text="Не получать вопросы на которые ответил", variable=self.var)
        self.rad4.place(x=165, y=188)
        
        
        self.answer = ttk.Button(self, text="Ответил", command=self.answerQ, state='disabled')
        self.answer.place(x=290, y=240)
        self.noAnswer = ttk.Button(self, text="Не Ответил", command=self.noAnswerQ, state='disabled')
        self.noAnswer.place(x=210, y=240)
        self.numAndQuestions = Message(self, padx=6, text=f'Тут будут появляться вопросы', bg='lightgrey', font="times 14 italic", width=550)
        self.numAndQuestions.grid(column=0, row=0)
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)

    def file(self):
        
        self.path = fd.askopenfilename()
        pathName = self.path
        self.name = self.path.split('/')[-1] if '/' in pathName else pathName.split()[-1]
        self.nameFile["text"] = "Файл: " + self.name
        docx = Document(pathName)
        countQuestions = len(docx.paragraphs)

        self.nextQues['state'] = 'disabled'
        self.beforeQues['state'] = 'disabled'
        self.answer['state'] = 'disabled'
        self.noAnswer['state'] = 'disabled'

        for i in range(0, countQuestions):
            
            textList = docx.paragraphs[i].text
            textList = list(textList.split())
            check = textList.count("|")
            ind = 0
            if check != 0:
                ind = textList.index("|")
                if textList[-2] == "Не" and textList[-1] == "Ответил":
                    self.dictQuestions[i+1] = [" ".join(textList[1:ind]), "Не Ответил"]
                    print("дошел до 2")
                elif textList[-2] == "|" and textList[-1] == "Ответил":
                    self.dictQuestions[i+1] = [" ".join(textList[1:ind]), "Ответил"]
                    print("дошел до 2")
                elif textList[-1] == "Затронул":
                    self.dictQuestions[i+1] = [" ".join(textList[1:ind]), "Не Затронул"]
                    print("дошел до 3")
            else:
                self.dictQuestions[i+1] = [docx.paragraphs[i].text, "Не Затронул"]


    def save_result(self):
            file_name = self.name
            save_filename = file_name.split('.')[0]
            directory = fd.askdirectory()
            print(directory+'/'+save_filename + ".docx")
            docx = Document()
            
            for i in range(1, len(self.dictQuestions.values())):
                text = f"{i}) {self.dictQuestions[i][0]} | "
                para = docx.add_paragraph(text)
                if self.dictQuestions[i][1] == "Ответил":
                    para.add_run(self.dictQuestions[i][1]).font.highlight_color = WD_COLOR_INDEX.GREEN
                elif self.dictQuestions[i][1] == "Не Ответил":
                    para.add_run(self.dictQuestions[i][1]).font.highlight_color = WD_COLOR_INDEX.RED
                else:
                    para.add_run(self.dictQuestions[i][1]).font.highlight_color = WD_COLOR_INDEX.GRAY_50
                docx.save(directory+'/'+save_filename + '.docx')
 

 
    def getNextQuest(self):
        try:
            self.questNumber = self.questNumber + 1
            if self.questNumber >= len(self.dictQuestions.values()):
                self.questNumber = 1
            text = f'Вопрос №{self.questNumber}:\n{self.dictQuestions[self.questNumber][0]}'
            print(text)
            self.numAndQuestions['text'] = text
            self.check_condition()
        except:
            self.numAndQuestions['text'] = "Сначала надо загрузить файл!"
            self.check_condition()
    
    def getBeforeQuest(self):
        try:
            self.questNumber = self.questNumber - 1
            if self.questNumber <= 0:
                self.questNumber = len(self.dictQuestions.values())
            
            text = f'Вопрос №{self.questNumber}:\n{self.dictQuestions[self.questNumber][0]}'
            print(text)

            self.numAndQuestions['text'] = text
            self.check_condition()
        except:
            self.numAndQuestions['text'] = "Сначала надо загрузить файл!"
            self.check_condition()


    def getQuestion(self):
        try:
            self.getQuestions['state'] = 'normal'
            
            flag = True
            while flag:
                number = random.randrange(1, len(self.dictQuestions.values()))
                
                if self.var.get() == 1:
                    print("Функция ВКЛЮЧЕНА")
                    if self.dictQuestions[number][1] != "Ответил" :
                        text = f'Вопрос №{number}:\n{self.dictQuestions[number][0]}'
                        break
                    else:
                        continue
                else:
                    print("Функция ВЫКЛЮЧЕНА")
                    text = f'Вопрос №{number}:\n{self.dictQuestions[number][0]}'
                    break
                    

            self.nextQues['state'] = 'normal'
            self.beforeQues['state'] = 'normal'
            self.answer['state'] = 'normal'
            self.noAnswer['state'] = 'normal'


            self.questNumber = number
            print(text)
            self.numAndQuestions['text'] = text
            self.check_condition()

            
        except ValueError:
            self.numAndQuestions['text'] = "Сначала надо загрузить файл!"
            self.numAndQuestions.configure(bg='red', font="times 14 italic", width=550)

    def answerQ(self):
        try:
            if self.dictQuestions[self.questNumber][1] != "Ответил" :
                self.dictQuestions[self.questNumber][1] = "Ответил" 
                self.check_condition()
            else:
                pass
        except ValueError:
            self.numAndQuestions['text'] = "Сначала надо загрузить файл!"
            self.numAndQuestions.configure(bg='red', font="times 14 italic", width=550)

    def noAnswerQ(self):
        try:
            if self.dictQuestions[self.questNumber][1] != "Не Ответил" :
                self.dictQuestions[self.questNumber][1] = "Не Ответил" 
                self.check_condition()
            else:
                pass
        except ValueError:
            self.numAndQuestions['text'] = "Сначала надо загрузить файл!"
            self.numAndQuestions.configure(bg='red', font="times 14 italic", width=550)

    def check_condition(self):
        """Проверка состояния вопроса и присвоние соответствующего цвета"""
        if self.dictQuestions[self.questNumber][1] == "Ответил":  
            self.numAndQuestions.configure(bg='lightgreen', font="times 14 italic", width=550)
        elif self.dictQuestions[self.questNumber][1] == "Не Ответил":  
            self.numAndQuestions.configure(bg='red', font="times 14 italic", width=550)
        elif self.dictQuestions[self.questNumber][1] == "Не Затронул":  
            self.numAndQuestions.configure(bg='lightgrey', font="times 14 italic", width=550)
    def confirm_delete(self):
        message = "Вы уверены, что хотите закрыть это окно?"
        if mb.askyesno(message=message, parent=self):
            self.destroy()


if __name__ == "__main__":
    app = MainWork()
    app.mainloop()

